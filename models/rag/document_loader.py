"""
文档加载器 - 加载和解析各种格式的文档
支持 PDF、Word、Excel、TXT 等格式
"""

import os
from typing import List, Dict, Optional
from dataclasses import dataclass


@dataclass
class Document:
    """文档数据结构"""
    content: str
    metadata: Dict
    id: Optional[str] = None

    def to_dict(self) -> Dict:
        return {
            'id': self.id,
            'content': self.content,
            'metadata': self.metadata
        }


class DocumentLoader:
    """文档加载器"""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def load_directory(self, directory: str, extensions: Optional[List[str]] = None) -> List[Document]:
        """
        加载目录下所有文档

        Args:
            directory: 文档目录路径
            extensions: 要加载的文件扩展名列表

        Returns:
            文档列表
        """
        if not os.path.exists(directory):
            raise FileNotFoundError(f"目录不存在：{directory}")

        if extensions is None:
            extensions = ['.pdf', '.doc', '.docx', '.txt', '.md', '.xlsx', '.xls']

        documents = []

        for root, dirs, files in os.walk(directory):
            for file in files:
                file_path = os.path.join(root, file)
                ext = os.path.splitext(file)[1].lower()

                if ext in extensions:
                    try:
                        docs = self.load_file(file_path)
                        documents.extend(docs)
                    except Exception as e:
                        print(f"加载文件失败 {file_path}: {e}")

        return documents

    def load_file(self, file_path: str) -> List[Document]:
        """加载单个文件"""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.pdf':
            return self._load_pdf(file_path)
        elif ext in ['.doc', '.docx']:
            return self._load_word(file_path)
        elif ext in ['.xlsx', '.xls']:
            return self._load_excel(file_path)
        elif ext == '.txt':
            return self._load_txt(file_path)
        elif ext == '.md':
            return self._load_markdown(file_path)
        else:
            return []

    def _load_pdf(self, file_path: str) -> List[Document]:
        """加载 PDF 文件"""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("请安装 PyMuPDF: pip install PyMuPDF")

        documents = []
        doc = fitz.open(file_path)

        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                doc = Document(
                    content=text,
                    metadata={
                        'source': file_path,
                        'page': page_num + 1,
                        'type': 'pdf'
                    },
                    id=f"{file_path}_page_{page_num + 1}"
                )
                documents.append(doc)

        doc.close()
        return documents

    def _load_word(self, file_path: str) -> List[Document]:
        """加载 Word 文件"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = DocxDocument(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        content = '\n'.join(text_parts)

        return [Document(
            content=content,
            metadata={
                'source': file_path,
                'type': 'word'
            },
            id=file_path
        )]

    def _load_excel(self, file_path: str) -> List[Document]:
        """加载 Excel 文件"""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("请安装 pandas: pip install pandas")

        documents = []
        excel_file = pd.ExcelFile(file_path)

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            # 将表格转换为文本
            text_parts = [f"表格：{sheet_name}"]
            text_parts.append(df.to_string())

            content = '\n'.join(text_parts)

            documents.append(Document(
                content=content,
                metadata={
                    'source': file_path,
                    'sheet': sheet_name,
                    'type': 'excel',
                    'rows': len(df),
                    'columns': list(df.columns)
                },
                id=f"{file_path}_{sheet_name}"
            ))

        return documents

    def _load_txt(self, file_path: str) -> List[Document]:
        """加载 TXT 文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        return [Document(
            content=content,
            metadata={
                'source': file_path,
                'type': 'txt'
            },
            id=file_path
        )]

    def _load_markdown(self, file_path: str) -> List[Document]:
        """加载 Markdown 文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        return [Document(
            content=content,
            metadata={
                'source': file_path,
                'type': 'md'
            },
            id=file_path
        )]

    def chunk_text(self, text: str) -> List[str]:
        """将长文本分块"""
        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]

            # 尝试在句子边界处切分
            if end < len(text):
                for sep in ['.\n', '。\n', '. ', '。', '\n\n', '\n']:
                    last_sep = chunk.rfind(sep)
                    if last_sep > self.chunk_size // 2:
                        chunk = text[start:start + last_sep + len(sep)]
                        break

            chunks.append(chunk.strip())
            start = end - self.chunk_overlap

        return chunks

    def load_and_chunk(self, file_path: str) -> List[Dict]:
        """加载文件并分块"""
        documents = self.load_file(file_path)
        chunks = []

        for doc in documents:
            if len(doc.content) > self.chunk_size:
                text_chunks = self.chunk_text(doc.content)
                for i, chunk in enumerate(text_chunks):
                    chunks.append({
                        'id': f"{doc.id}_chunk_{i}",
                        'content': chunk,
                        'metadata': {
                            **doc.metadata,
                            'chunk_index': i,
                            'total_chunks': len(text_chunks)
                        }
                    })
            else:
                chunks.append({
                    'id': doc.id,
                    'content': doc.content,
                    'metadata': doc.metadata
                })

        return chunks
